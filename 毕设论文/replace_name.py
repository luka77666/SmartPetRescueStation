import os
import shutil
from docx import Document

root = r"c:\Users\86157\Desktop\ronghe"
backup_dir = os.path.join(root, "_backup_罗志杰改罗智杰")
os.makedirs(backup_dir, exist_ok=True)

replace_count = 0  # 替换文件数
total_names = 0    # 替换名字数

def replace_in_paragraphs(paragraphs):
    """替换段落中的文字，返回替换次数"""
    cnt = 0
    for p in paragraphs:
        if '罗志杰' in p.text:
            for run in p.runs:
                if '罗志杰' in run.text:
                    run.text = run.text.replace('罗志杰', '罗智杰')
                    cnt += 1
    return cnt

def replace_in_tables(tables):
    """替换表格中的文字"""
    cnt = 0
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                cnt += replace_in_paragraphs(cell.paragraphs)
    return cnt

for dirpath, dirnames, filenames in os.walk(root):
    for fn in filenames:
        if fn.lower().endswith(('.docx',)):
            fp = os.path.join(dirpath, fn)
            try:
                doc = Document(fp)
                # 检查是否包含目标文字
                need_replace = False
                for p in doc.paragraphs:
                    if '罗志杰' in p.text:
                        need_replace = True
                        break
                if not need_replace:
                    for t in doc.tables:
                        for row in t.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    if '罗志杰' in p.text:
                                        need_replace = True
                                        break
                                if need_replace:
                                    break
                        if need_replace:
                            break
                if not need_replace:
                    continue

                # 备份原文件
                rel = os.path.relpath(fp, root)
                backup_sub = os.path.join(backup_dir, os.path.dirname(rel))
                os.makedirs(backup_sub, exist_ok=True)
                shutil.copy2(fp, os.path.join(backup_sub, os.path.basename(rel)))
                print("[备份] {}".format(rel))

                # 执行替换
                cnt = 0
                cnt += replace_in_paragraphs(doc.paragraphs)
                cnt += replace_in_tables(doc.tables)

                if cnt > 0:
                    doc.save(fp)
                    replace_count += 1
                    total_names += cnt
                    print("[替换] {} 处：{}".format(cnt, rel))
                else:
                    print("[跳过] 内容不匹配（表格/页眉等）：{}".format(rel))

            except Exception as e:
                print("[错误] {}: {}".format(fp, e))

print("=" * 60)
print("完成！共替换 {} 个文件，{} 处「罗志杰」→「罗智杰」".format(replace_count, total_names))
print("原文件备份在：{}".format(backup_dir))
print("=" * 60)
