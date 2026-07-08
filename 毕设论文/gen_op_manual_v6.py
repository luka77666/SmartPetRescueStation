# -*- coding: utf-8 -*-
"""按参考模板格式生成操作手册 V6 — 去掉小黑点，增加图片位置"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, datetime

OUTPUT = r'c:\Users\86157\Desktop\ronghe\毕设论文\基于物联网的智能宠物喂养控制系统V1.0操作手册_{}.docx'.format(
    datetime.datetime.now().strftime('%Y%m%d_%H%M%S'))

doc = Document()

# ========== 页面 A4 ==========
sec = doc.sections[0]
sec.page_width   = Cm(21)
sec.page_height  = Cm(29.7)
sec.top_margin   = Cm(2.5)
sec.bottom_margin= Cm(2.5)
sec.left_margin  = Cm(2.5)
sec.right_margin = Cm(2.5)

# ========== 配置 Normal 样式（正文基准）==========
normal_style = doc.styles['Normal']
normal_style.font.name = '宋体'
normal_style.font.size = Pt(11)
normal_style.paragraph_format.line_spacing = 1.5
normal_style.paragraph_format.first_line_indent = Cm(0.74)
normal_style.paragraph_format.space_after = Pt(6)
nf = normal_style.element.get_or_add_rPr()
nf_rfonts = OxmlElement('w:rFonts')
nf_rfonts.set(qn('w:eastAsia'), '宋体')
nf.append(nf_rfonts)

# ========== 配置 Heading 1 样式 ==========
h1_style = doc.styles['Heading 1']
h1_style.font.name = '黑体'
h1_style.font.size = Pt(15)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0, 0, 0)
pf = h1_style.paragraph_format
pf.space_before = Pt(24)
pf.space_after  = Pt(8)
pf.line_spacing = 1.5
pf.first_line_indent = Cm(0)
hf = h1_style.element.get_or_add_rPr()
hf_rfonts = OxmlElement('w:rFonts')
hf_rfonts.set(qn('w:eastAsia'), '黑体')
hf.append(hf_rfonts)

# ========== 配置 Heading 2 样式 ==========
h2_style = doc.styles['Heading 2']
h2_style.font.name = '黑体'
h2_style.font.size = Pt(14)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0, 0, 0)
pf2 = h2_style.paragraph_format
pf2.space_before = Pt(12)
pf2.space_after  = Pt(6)
pf2.line_spacing = 1.5
pf2.first_line_indent = Cm(0)
hf2 = h2_style.element.get_or_add_rPr()
hf2_rfonts = OxmlElement('w:rFonts')
hf2_rfonts.set(qn('w:eastAsia'), '黑体')
hf2.append(hf2_rfonts)

# ========== 辅助函数 ==========
def set_cjk(run, name):
    """设置东亚字体"""
    try:
        rPr = run._r.get_or_add_rPr()
        rfonts = rPr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rPr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), name)
    except:
        pass

def p(text, style=None, align='left', sb=None, sa=None, indent=True, bold=None, size=None):
    """添加一个段落"""
    if style:
        p_obj = doc.add_paragraph(text, style=style)
    else:
        p_obj = doc.add_paragraph(text)
    pf = p_obj.paragraph_format
    if align == 'center':
        p_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p_obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = 1.5
    if sb is not None:
        pf.space_before = Pt(sb)
    if sa is not None:
        pf.space_after = Pt(sa)
    if indent and align == 'left' and not style:
        pf.first_line_indent = Cm(0.74)
    elif not indent:
        pf.first_line_indent = Cm(0)
    for run in p_obj.runs:
        if size:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        set_cjk(run, '宋体')
    if not p_obj.runs and text:
        r = p_obj.add_run(text)
        if size:
            r.font.size = Pt(size)
        if bold is not None:
            r.font.bold = bold
        set_cjk(r, '宋体')
    return p_obj

def fig_placeholder(caption_text):
    """图片占位符 + 图注"""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.line_spacing = 1.5
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run('【此处插入图片：' + caption_text + '】')
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    set_cjk(r1, '宋体')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(caption_text)
    r2.font.size = Pt(11)
    r2.font.bold = True
    set_cjk(r2, '宋体')
    return p1, p2

def page_break():
    doc.add_page_break()

# ============================================================
#  封面
# ============================================================
p('', sa=30)
p('基于物联网的智能宠物喂养控制系统 V1.0',
  align='center', sb=40, sa=8, indent=False, bold=True, size=24)
p('操  作  手  册',
  align='center', sb=8, sa=40, indent=False, bold=True, size=26)

info_items = [
    ('著作权人：', '林湘龙'),
    ('开发完成日期：', '2026年5月'),
    ('首次发表日期：', '未发表'),
    ('版本号：', 'V1.0'),
]
for label, val in info_items:
    p_obj = doc.add_paragraph()
    p_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_obj.paragraph_format.left_indent = Cm(3.0)
    p_obj.paragraph_format.line_spacing = 1.5
    p_obj.paragraph_format.space_after = Pt(6)
    r1 = p_obj.add_run(label)
    r1.font.size = Pt(15)
    r1.font.bold = True
    set_cjk(r1, '黑体')
    r2 = p_obj.add_run(val)
    r2.font.size = Pt(14)
    set_cjk(r2, '宋体')

p('', sa=60)
p('仲恺农业工程学院', align='center', sb=30, sa=8, indent=False, bold=True, size=15)
p('2026年6月', align='center', sb=0, sa=40, indent=False, bold=True, size=14)

# ============================================================
#  目录
# ============================================================
page_break()
p('目    录', align='center', sb=10, sa=20, indent=False, bold=True, size=26)

toc = [
    ('第1章  系统概述',                  0),
    ('    1.1  系统主要功能',           1),
    ('    1.2  系统架构说明',           1),
    ('    1.3  适用对象',               1),
    ('第2章  运行环境',                  0),
    ('    2.1  软件环境',               1),
    ('    2.2  硬件环境',               1),
    ('第3章  安装与启动',                0),
    ('    3.1  APP安装方法',            1),
    ('    3.2  设备连接配置',            1),
    ('    3.3  首次启动与初始化',        1),
    ('第4章  操作详解',                  0),
    ('    4.1  首页数据监控',            1),
    ('    4.2  喂食控制操作',            1),
    ('    4.3  定时设置操作',            1),
    ('    4.4  阈值设置操作',            1),
    ('第5章  数据通信协议说明',           0),
    ('第6章  常见问题与维护',             0),
    ('附录A  术语表',                    0),
]
for text, level in toc:
    p_obj = doc.add_paragraph()
    pf = p_obj.paragraph_format
    pf.line_spacing = 1.8
    pf.space_after = Pt(4)
    if level > 0:
        pf.left_indent = Cm(0.75 * level)
    r1 = p_obj.add_run(text + '  ……  ')
    r1.font.size = Pt(12) if level == 0 else Pt(11)
    r1.font.bold = (level == 0)
    set_cjk(r1, '宋体')
    p_obj.paragraph_format.first_line_indent = Cm(0)
    doc.add_paragraph()

# ============================================================
#  第1章 系统概述
# ============================================================
page_break()
p('第1章  系统概述', style='Heading 1')

p_obj = p('基于物联网的智能宠物喂养控制系统',
          align='center', sb=8, sa=8, indent=False, bold=True, size=16)

p('基于物联网的智能宠物喂养控制系统V1.0（以下简称"本软件"）是一款面向宠物饲养场景的'
  'Android平台智能控制应用软件。本软件与智能宠物救助站硬件设备配套使用，'
  '通过WiFi无线通信技术实现手机与设备之间的数据交互，使用户能够随时随地通过手机对宠物饲养环境进行远程监控与智能管理。')

p('本软件采用Kotlin语言开发，基于Jetpack Compose现代化UI框架构建用户界面，'
  '使用Navigation Compose组件实现多页面导航管理。软件通过TCP Socket通信协议与设备端ESP8266 WiFi模块建立连接，'
  '实现传感器数据的实时采集展示、喂养指令的远程下发、定时任务的配置管理等功能，'
  '有效解决了宠物主人外出期间宠物无人照料的实际问题。')

# 1.1
p('1.1  系统主要功能', style='Heading 2')
for item in [
    '实时环境监测：通过APP首页实时显示设备上传的温度、湿度、重量、水位四项环境参数，数据每5秒自动刷新一次。',
    '自动喂食控制：支持自动模式和手动模式两种控制方式，自动模式下系统根据预设称重阈值自动触发喂食，手动模式下用户通过APP按钮控制喂食启动与停止。',
    '自动补水控制：系统实时监测水位百分比，当水位低于设定阈值时自动启动水泵进行补水，到达阈值后自动停止。',
    'APP远程监控：用户通过Android手机安装本软件后，连接设备WiFi热点即可实现远程监控，无需路由器或互联网连接。',
    '定时喂食计划：支持3组独立定时任务，每组可设置喂食时间和目标喂食量，并支持单独启用或禁用，满足不同时段的喂养需求。',
    '阈值参数设置：支持通过APP或设备按键设置水位阈值（0~100%）和称重阈值（0~9999，单位0.1g），参数自动保存至设备Flash，断电不丢失。',
    '设备状态显示：APP实时显示设备运行模式（自动/手动）、喂食状态（喂食中/待机）、水泵状态（开/关），以及各项参数的当前设定值。',
    '时间自动同步：APP连接设备成功后自动将手机时间同步至设备RTC实时时钟，确保定时任务准确执行。',
]:
    p(item)

# 1.2
p('1.2  系统架构说明', style='Heading 2')
p('本软件采用客户端-服务端（C/S）架构设计。Android APP作为客户端，通过TCP协议连接设备端ESP8266 WiFi模块（AP模式，TCP Server端口5000）。'
  '设备端作为服务端，负责采集传感器数据、执行控制指令，并将运行状态实时反馈至APP端。'
  '通信数据采用自定义文本协议，以换行符作为帧结束标志，数据格式紧凑高效，解析速度快。')

fig_placeholder('图1-1  系统总体架构图')
p('（请在Word中将上方文字替换为实际系统架构截图。）')

# 1.3
p('1.3  适用对象', style='Heading 2')
for item in [
    '宠物主人：外出工作、旅行或临时离家时，通过本软件远程监控宠物饲养环境，定时定量喂养宠物。',
    '宠物寄养机构：对多只宠物进行规范化、定时化喂养管理，提升寄养服务质量。',
    '宠物救助站：为流浪动物提供自动化喂养服务，降低人工成本，提高救助效率。',
]:
    p(item)

# ============================================================
#  第2章 运行环境
# ============================================================
page_break()
p('第2章  运行环境', style='Heading 1')

p('2.1  软件环境', style='Heading 2')
p('本软件运行于Android操作系统，具体软件环境要求如下：')
for item in [
    '操作系统：Android 8.0（API Level 26）及以上版本。',
    '开发语言：Kotlin（编程语言），Jetpack Compose（UI框架）。',
    '开发工具：Android Studio（Hedgehog | 2023.1.1 或更高版本）。',
    '编译环境：Gradle 8.2，Kotlin 1.9.22，Compose BOM 2024.02.00。',
    '运行依赖：AndroidX Compose组件库、Navigation Compose导航组件、DataStore Preferences数据存储。',
    '网络通信：Android系统TCP/IP协议栈，支持WiFi无线网络连接。',
]:
    p(item)

p('2.2  硬件环境', style='Heading 2')
p('本软件需要与智能宠物救助站硬件设备配套使用，对Android终端设备的最低硬件要求如下：')
for item in [
    'CPU：四核1.5GHz及以上（ARMv8架构优先）。',
    'RAM：2GB及以上（建议4GB以上以确保流畅运行）。',
    '存储空间：安装包约15MB，运行时数据占用不超过50MB。',
    '显示屏幕：分辨率不低于720×1280像素，建议1080×2400及以上。',
    'WiFi模块：支持802.11 b/g/n无线协议，能够连接设备AP热点。',
    '操作系统版本：Android 8.0及以上，建议Android 10.0以上以获得最佳体验。',
]:
    p(item)
p('配套硬件设备（智能宠物救助站）主要参数：'
  '主控芯片STM32F103C8T6（Cortex-M3内核，72MHz主频，64KB Flash，20KB RAM）；'
  'WiFi通信模块ESP8266-01S（AP模式，SSID：ZNCWJZZ，TCP Server端口5000）；'
  '传感器单元包括DHT11温湿度传感器（单总线接口）、HX711称重传感器（24位ADC，精度0.1g）、水位传感器（ADC模拟量输入）；'
  '执行机构包括28BYJ-48步进电机（ULN2003驱动）、5V微型水泵（继电器控制）。')

# ============================================================
#  第3章 安装与启动
# ============================================================
page_break()
p('第3章  安装与启动', style='Heading 1')

p('3.1  APP安装方法', style='Heading 2')
p('本软件以APK安装包形式提供给用户，安装步骤如下：')
for i, item in enumerate([
    '获取安装包：将APP安装包（文件名为 app-release.apk）通过数据线、蓝牙或网络传输方式复制到Android手机存储中。',
    '开启安装权限：在Android手机"设置"→"安全"→"更多安全设置"中，开启"允许安装未知来源应用"选项（不同品牌手机设置路径可能略有差异）。',
    '执行安装：使用手机文件管理器找到app-release.apk文件，点击后按系统提示完成安装。安装过程中如出现"该应用可能存在风险"提示，选择"继续安装"即可。',
    '安装完成：安装成功后，手机桌面将出现"智能宠物救助站"应用图标，点击图标即可启动软件。',
], 1):
    p('（{}）{}'.format(i, item))
fig_placeholder('图3-1  APP安装步骤示意图')

p('3.2  设备连接配置', style='Heading 2')
p('本软件通过WiFi与智能宠物救助站硬件设备建立连接，设备连接配置步骤如下：')
for i, item in enumerate([
    '设备通电：将智能宠物救助站接通电源（DC 5V/2A电源适配器），设备自动启动，ESP8266模块约5~8秒后建立AP热点。',
    '连接设备热点：打开手机"设置"→"WiFi"，在可用网络列表中找到名称为"ZNCWJZZ"的WiFi热点，点击连接（该热点为开放网络，无需密码）。',
    '启动APP：点击手机桌面"智能宠物救助站"应用图标，启动本软件。',
    '输入设备IP地址：在APP首页顶部的IP地址输入框中，输入设备IP地址"192.168.4.1"，点击"连接"按钮。',
    '连接成功：连接成功后，APP自动跳转至主页界面，开始接收设备上传的传感器数据，连接状态显示为"已连接"。',
], 1):
    p('（{}）{}'.format(i, item))
fig_placeholder('图3-2  设备连接配置示意图')
p('注意：如连接失败，请确认手机已正确连接"ZNCWJZZ"热点，且IP地址输入正确。'
  '如仍无法连接，请重启设备后重试。APP支持自动重连功能，网络断开后将自动尝试重新连接。')

p('3.3  首次启动与初始化', style='Heading 2')
p('首次启动APP后，软件将自动执行以下初始化操作：')
for item in [
    '自动检测设备WiFi热点连接状态，如未连接将提示用户先连接设备热点。',
    '连接成功后自动发送TIME命令，将手机系统时间同步至设备RTC实时时钟。',
    '自动请求设备当前传感器数据和运行状态，刷新首页显示。',
    '加载上次保存的连接配置（如有），简化下次连接流程。',
]:
    p(item)
fig_placeholder('图3-3  APP首次启动与初始化界面')

# ============================================================
#  第4章 操作详解
# ============================================================
page_break()
p('第4章  操作详解', style='Heading 1')
p_obj = p('基于物联网的智能宠物喂养控制系统',
          align='center', sb=8, sa=8, indent=False, bold=True, size=16)

# 4.1
p('4.1  首页数据监控', style='Heading 2')
p('首页（监测页面）是APP启动后的默认界面，用于实时展示设备运行状态和环境参数。')
for item in [
    '页面布局：顶部为校徽图标和系统名称；中部为四个数据卡片，分别显示温度（℃）、湿度（%RH）、重量（g）、水位（%）；底部为三个导航标签（首页/喂食/定时）。',
    '数据刷新：APP每5秒向设备发送一次数据请求，设备收到请求后返回当前传感器数据，APP解析后在对应卡片中更新显示。',
    '报警提示：当检测数值超出设定阈值时（如水位低于设定阈值），对应数据卡片的数值显示为红色，提醒用户及时关注。',
    '连接状态：页面顶部显示当前连接状态（"已连接"或"未连接"），未连接时数据卡片显示"--"占位符。',
    '时间同步：APP连接设备成功后，自动将手机系统时间通过TIME命令同步至设备RTC实时时钟，确保定时任务准确执行。',
]:
    p(item)
fig_placeholder('图4-1  APP首页数据监控界面')

p('首页报警状态说明：当设备运行异常或环境参数超出安全范围时，APP会以红色字体进行报警提示。')
for item in [
    '温度报警：当环境温度超出设定安全范围时，温度卡片数值显示为红色，提醒用户注意宠物饲养环境温度。',
    '水位报警：当水位百分比低于设定阈值时，水位卡片数值显示为红色，提示用户及时补水或检查水泵。',
    '重量报警：当称重数值异常（如传感器故障或饲料空仓）时，重量卡片数值显示为红色，提醒用户检查设备。',
]:
    p(item)
fig_placeholder('图4-2  首页报警提示状态界面')

p('连接断开状态：当设备断电、WiFi信号中断或APP长时间未收到设备数据时，首页顶部连接状态显示为"未连接"，'
  '四个数据卡片数值显示为"--"，此时用户需检查设备供电和WiFi连接后重新点击连接按钮。')
fig_placeholder('图4-3  首页连接断开状态界面')

# 4.2
p('4.2  喂食控制操作', style='Heading 2')
p('喂食控制页面用于手动控制喂食动作、切换设备运行模式、查看设备工作状态。'
  '在底部导航栏点击"喂食"标签即可进入该页面。')
p('页面布局：页面顶部显示当前称重数值和水位数值；中部显示设备运行模式（自动/手动）、喂食状态（喂食中/待机）、水泵状态（开/关）；底部为四个操作按钮（开始喂食、停止喂食、切换模式、进入设置）。')
p('开始喂食：点击"开始喂食"按钮，APP先发送KB命令将光标移动至喂食控制位，再发送KC命令启动喂食，步进电机正转推动饲料下落，同时APP显示状态更新为"喂食中"。')
p('停止喂食：点击"停止喂食"按钮，APP发送KB+KD命令停止电机转动，步进电机停止转动，APP显示状态更新为"待机"。')
p('切换模式：点击"切换模式"按钮，APP发送KA命令，设备在自动模式和手动模式之间切换，APP界面同步更新模式显示。自动模式下系统根据称重阈值自动触发喂食，手动模式下需用户手动控制。')

p('阈值调节', style='Heading 3')
p('增量称重：手动喂食过程中，系统采用增量称重方式判断停止条件。记录开始喂食前的初始重量，当当前重量大于等于初始重量加目标喂食量时自动停止喂食，确保喂食精度。该机制可有效消除食盆皮重变化对称重精度的影响，避免因食盆内剩余饲料重量变化导致喂食量偏差。')
p('防卡粮机制：喂食过程中如发生饲料卡住导致电机堵转的情况，系统自动触发防卡粮程序——正转5圈后反转1圈，循环执行，直至称重达到目标值或用户手动停止。该机制有效防止因饲料潮湿结块或颗粒过大导致的电机堵转问题，提高设备运行可靠性。')
fig_placeholder('图4-4  APP喂食控制界面')

p('喂食过程状态展示：开始喂食后，APP界面实时更新以下状态信息：')
for item in [
    '喂食状态标签从"待机"变为"喂食中"，以醒目的颜色标识当前正在执行喂食动作。',
    '称重数值实时变化，用户可观察饲料重量增加情况，直观了解喂食进度。',
    '当称重增量达到设定阈值时，系统自动停止喂食，状态标签恢复为"待机"，无需用户手动干预。',
]:
    p(item)
fig_placeholder('图4-5  喂食过程中状态界面')

p('停止喂食与待机状态：喂食完成后或用户主动点击"停止喂食"后，系统进入待机状态：')
for item in [
    '喂食状态标签显示为"待机"，表示当前无喂食动作执行。',
    '称重数值保持在当前测量值，作为下次喂食的参考基准。',
    '用户可随时再次点击"开始喂食"启动新的喂食流程。',
]:
    p(item)
fig_placeholder('图4-6  喂食停止后待机状态界面')

p('自动模式与手动模式切换：设备支持两种运行模式，用户可根据实际需求灵活切换：')
for item in [
    '自动模式：系统根据预设的称重阈值和水位阈值自动判断是否需要喂食或补水，无需人工干预，适合日常规律性喂养场景。',
    '手动模式：所有喂食和补水动作由用户通过APP按钮或设备按键手动控制，适合临时调整喂养计划或测试设备功能。',
    '模式切换：点击"切换模式"按钮，APP发送KA命令，设备立即切换模式，APP界面同步更新模式显示。',
]:
    p(item)
fig_placeholder('图4-7  自动模式与手动模式切换界面')

p('水泵控制操作：在喂食控制页面，用户可查看当前水泵状态，并根据需要手动控制水泵：')
for item in [
    '查看水泵状态：页面中部显示"水泵状态"标签，显示当前水泵为"开"或"关"。',
    '手动开启水泵：在设备端将光标移动至水泵控制位，按KEY_3（或APP发送KC命令）启动水泵。',
    '手动关闭水泵：按KEY_4（或APP发送KD命令）关闭水泵。',
    '自动补水：自动模式下，当水位低于设定阈值时，系统自动启动水泵进行补水，到达阈值后自动停止。',
]:
    p(item)
fig_placeholder('图4-8  水泵开启与关闭状态界面')

# 4.3
p('4.3  定时设置操作', style='Heading 2')
p('定时设置页面用于配置3组独立定时喂食任务，支持分别设置喂食时间、喂食量和启用状态。'
  '在底部导航栏点击"定时"标签即可进入该页面。')
for item in [
    '页面布局：页面以卡片形式展示3组定时任务（定时1/定时2/定时3），每组卡片包含启用开关（Switch）、喂食时间（时/分滚轮选择器）、喂食量（数字输入框，单位g）。页面底部为"全部保存到设备"按钮。',
    '设置喂食时间：点击时间滚轮区域，小时滚轮和分钟滚轮分别通过上下滑动选择，选中项以蓝色高亮显示，非选中项逐渐淡化，操作体验类似手机闹钟时间设置。',
    '设置喂食量：在对应定时卡片的喂食量输入框中输入目标重量数值（单位：g，建议设置范围为10~200g），系统将根据此数值控制喂食停止条件。',
    '启用/禁用定时：点击定时卡片右上角的启用开关（Switch），开关为开启状态时该组定时任务生效，为关闭状态时该组定时任务暂停执行。',
    '保存到设备：完成所有设置后，点击页面底部"全部保存到设备"按钮，APP将依次发送9条命令至设备（3组时间的STn命令+3组喂食量的SDn命令+3组启用状态的SEn命令），每条命令间隔100ms，确保设备正确接收。',
    '定时执行逻辑：设备端RTC实时时钟到达设定时间后，自动触发喂食流程，控制步进电机转动，当称重增量达到设定喂食量时自动停止，完成一次定时喂食任务。',
]:
    p(item)
fig_placeholder('图4-9  APP定时设置总览界面')

p('定时1设置详解：定时1为第一组定时任务，通常用于设置早晨喂食时间。')
for item in [
    '喂食时间设置：通过时/分滚轮选择器设置定时1的触发时间，例如设置07:00表示每天早晨7点自动执行喂食。',
    '喂食量设置：在喂食量输入框中输入目标重量，例如输入50表示喂食50g饲料后自动停止。',
    '启用开关：打开启用开关后，定时1任务生效；关闭后该任务暂停执行，但不删除已设置的时间和喂食量。',
]:
    p(item)
fig_placeholder('图4-10  定时1设置特写界面')

p('定时2设置详解：定时2为第二组定时任务，通常用于设置中午或下午喂食时间。')
for item in [
    '喂食时间设置：例如设置12:30表示每天中午12点30分自动执行喂食。',
    '喂食量设置：根据宠物食量设置合适的喂食量，建议与定时1保持相同或略有差异。',
    '启用开关：独立控制定时2的启用状态，与定时1、定时3互不干扰。',
]:
    p(item)
fig_placeholder('图4-11  定时2设置特写界面')

p('定时3设置详解：定时3为第三组定时任务，通常用于设置傍晚或夜间喂食时间。')
for item in [
    '喂食时间设置：例如设置19:00表示每天晚上7点自动执行喂食。',
    '喂食量设置：建议晚间喂食量适当减少，避免宠物夜间积食。',
    '启用开关：独立控制定时3的启用状态，三组定时可同时启用，也可根据需要选择性启用。',
]:
    p(item)
fig_placeholder('图4-12  定时3设置特写界面')

p('全部保存与执行确认：完成所有定时设置后，点击"全部保存到设备"按钮，APP将执行以下操作：')
for item in [
    '依次发送ST1、ST2、ST3命令，将三组定时时间写入设备Flash存储器。',
    '依次发送SD1、SD2、SD3命令，将三组喂食量写入设备Flash存储器。',
    '依次发送SE1、SE2、SE3命令，将三组启用状态写入设备Flash存储器。',
    '每条命令间隔100ms，确保设备有足够时间处理和响应。',
    '保存成功后，设备将在到达设定时间时自动执行对应的定时喂食任务。',
]:
    p(item)
fig_placeholder('图4-13  全部保存成功提示界面')

# 4.4
p('4.4  阈值设置操作', style='Heading 2')
p('阈值设置用于调整设备自动模式下的触发条件，包括水位阈值和称重阈值。'
  '阈值可通过设备端按键设置，也可通过APP发送命令设置。')
for item in [
    '水位阈值：设置自动补水的触发水位百分比（0~100%）。当设备检测到水位低于此阈值时，自动启动水泵进行补水；当水位达到此阈值时，自动停止水泵。建议设置值为80%。',
    '称重阈值（手动喂食停止量）：设置手动模式下单次喂食的目标重量（单位：g）。点击"开始喂食"后，系统持续监测称重数值，当增量达到此阈值时自动停止喂食。建议设置值为50g。',
    'APP命令设置方式：在APP喂食页面点击"进入设置"后，通过发送SWxxx命令（设置水位阈值）和SFxxxx命令（设置称重阈值）进行设置，设备收到命令后自动保存至Flash存储器。',
    '设备按键设置方式：在设备端长按KEY_1进入设置页面，短按KEY_2移动光标至对应参数，通过KEY_3/KEY_4调整数值，设置完成后等待10秒自动退出并保存。',
    '参数掉电保存：所有阈值参数均保存至STM32内部Flash存储器（地址0x0800F000），设备断电后参数不丢失，下次上电自动加载上次设定值。',
]:
    p(item)
fig_placeholder('图4-14  阈值设置界面（水位阈值与称重阈值）')

p('阈值设置注意事项：')
for item in [
    '水位阈值不宜设置过高（如95%以上），否则可能导致水泵频繁启停，影响设备寿命。',
    '称重阈值应根据宠物实际食量设置，过小可能导致喂食不足，过大可能导致饲料浪费。',
    '修改阈值后建议观察1~2个运行周期，确认设备按新阈值正常工作后再离开。',
]:
    p(item)
fig_placeholder('图4-15  阈值设置成功后的设备运行界面')

# ============================================================
#  第5章 数据通信协议说明
# ============================================================
page_break()
p('第5章  数据通信协议说明', style='Heading 1')
p_obj = p('基于物联网的智能宠物喂养控制系统',
          align='center', sb=8, sa=8, indent=False, bold=True, size=16)

p('本软件与智能宠物救助站硬件设备之间通过TCP Socket协议进行数据通信，'
  '通信方式为：APP作为TCP客户端，设备端ESP8266作为TCP服务端（端口5000）。'
  '本章对通信协议格式进行详细说明，便于软件维护和功能扩展时参考。')

p('5.1  数据传输格式', style='Heading 2')
p('APP与设备之间的数据传输采用文本字符串格式，每行以"\\r\\n"作为结束标志。'
  'APP向设备发送命令为单行字符串，设备向APP返回数据为单行字符串。')

p('5.2  设备上传数据帧格式', style='Heading 2')
p('设备每隔1秒主动上传一次传感器数据，数据帧格式如下：')
p('T25H60W0200L080M1F0P0WT080FT0200T10700F10050E1T21230F20100E1T31900F30150E1')
p('各字段含义说明：'
  'T=温度（℃），H=湿度（%RH），W=重量（g），L=水位（%），'
  'M=运行模式（1=自动/0=手动），F=喂食状态（1=喂食中/0=待机），P=水泵状态（1=开/0=关），'
  'WT=水位阈值，FT=称重阈值，'
  'T1/T2/T3=定时时间（时分格式），F1/F2/F3=定时喂食量（g），E1/E2/E3=定时启用状态（1=启用/0=禁用）。')
fig_placeholder('图5-1  数据通信帧格式解析示意图')

p('5.3  APP下发命令格式', style='Heading 2')
for item in [
    'KA：切换自动/手动模式（发送后设备模式翻转）。',
    'KB：光标移动键，在设置页面中循环切换光标位置（1=水泵/2=喂食/3=水位阈值/4=称重阈值/5=定时1/6=定时2/7=定时3）。',
    'KC：确认/开启键，根据当前光标位置执行对应开启操作（光标在喂食位时启动喂食，在水泵位时启动水泵）。',
    'KD：关闭/返回键，根据当前光标位置执行对应关闭操作（光标在喂食位时停止喂食，在水泵位时关闭水泵）。',
    'KE：长按KEY_1对应命令，进入设备设置页面。',
    'SWxxx：设置水位阈值（xxx为三位数字，范围000~100，例如SW080表示设置水位阈值为80%）。',
    'SFxxxx：设置称重阈值（xxxx为四位数字，范围0000~9999，单位0.1g，例如SF0050表示设置称重阈值为50g）。',
    'STnHHMM：设置定时n的时间（n=1/2/3，HHMM为时分，例如ST10700表示定时1设置为07:00）。',
    'SDnXXXX：设置定时n的喂食量（n=1/2/3，XXXX为四位数，单位g，例如SD10500表示定时1喂食量设置为500g）。',
    'SEn0/SEn1：设置定时n的启用状态（0=禁用/1=启用，例如SE11表示启用定时1）。',
    'TIMEHHMMSS：时间同步命令，将手机时间写入设备RTC（例如TIME073500表示同步时间为07:35:00）。',
    'MF0/MF1：直通命令，直接启动/停止喂食，不依赖光标位置。',
    'MW0/MW1：直通命令，直接开启/关闭水泵，不依赖光标位置。',
]:
    p(item)

p('5.4  心跳保活机制', style='Heading 2')
p('为保持TCP连接稳定性，APP每5秒向设备发送一次"PING"心跳包，设备收到后回复"PONG"。'
  '如连续3次未收到"PONG"回复，APP判定连接已断开，自动触发重连机制。')
fig_placeholder('图5-2  心跳保活机制时序图')

# ============================================================
#  第6章 常见问题与维护
# ============================================================
page_break()
p('第6章  常见问题与维护', style='Heading 1')
p_obj = p('基于物联网的智能宠物喂养控制系统',
          align='center', sb=8, sa=8, indent=False, bold=True, size=16)

p('6.1  常见问题', style='Heading 2')
p('以下列出本软件使用过程中可能遇到的常见问题及解决方法：')

faq = [
    ('APP无法连接设备',
     '请按以下步骤排查：①确认手机已连接设备热点"ZNCWJZZ"；②确认APP中IP地址填写为"192.168.4.1"；③确认端口号为5000（APP内部已固定，用户无需手动填写）；④如仍无法连接，重启设备后重试；⑤检查设备是否正常上电，ESP8266模块指示灯是否正常闪烁。'),
    ('首页数据不更新或显示"--"',
     '此现象通常为连接断开所致。请检查WiFi连接是否正常，或点击首页"连接"按钮重新连接。如连接状态显示"已连接"但数据仍不更新，请退出APP后重新进入。'),
    ('称重数值不准或偏差较大',
     '请先断电重启设备进行自动归零（去皮）。如仍有偏差，检查HX711称重传感器是否受压变形或接线是否牢固。本系统称重范围为0~500g，超出范围可能导致读数不准。'),
    ('点击"开始喂食"后电机不转',
     '请按以下步骤排查：①确认设备处于手动模式（如在自动模式下，点击"切换模式"按钮切换）；②确认已通过KEY_2（或APP发送KB命令）将光标移动至"喂食"位置，再按KEY_3（或APP发送KC命令）启动；③检查步进电机接线及ULN2003驱动板供电是否正常；④检查饲料是否卡住，如卡住系统会自动触发防卡粮程序。'),
    ('水泵不工作或不自动补水',
     '请检查：①水位传感器连接是否正常；②水位阈值设置是否正确（建议设置为80%）；③水泵供电及继电器控制模块是否正常；④在APP喂食页面查看"水泵状态"显示，如显示"关"可尝试手动点击"开启水泵"；⑤水位传感器ADC量程是否在28~4068范围内。'),
    ('定时喂食未按时触发',
     '请检查：①定时任务的启用开关是否已打开（Switch为开启状态）；②设备时间与设定时间是否一致（连接APP后时间会自动同步，如未连接APP请通过设备按键手动设置时间）；③定时时间格式为24小时制，请注意上午/下午时间设置是否正确。'),
    ('APP闪退或卡死',
     '请尝试：①退出APP后重新启动；②重启手机后重试；③检查手机Android系统版本是否在8.0以上；④如问题持续存在，请卸载后重新安装最新版本APK。'),
    ('设备端屏幕显示异常或乱码',
     '请断电重启设备，如重启后问题仍存在，检查TFT LCD排线是否松动，或联系技术人员进行维修。'),
]
for q, a in faq:
    p('问题：' + q, bold=True, sb=8, sa=4, indent=False)
    p('答：' + a)

fig_placeholder('图6-1  故障排查流程图')

p('6.2  日常维护', style='Heading 2')
for item in [
    '定期清洁：建议每周清洁一次饲料仓和称重传感器，避免饲料残渣堆积影响称重精度。',
    '水位传感器维护：每月检查一次水位传感器表面是否附着污垢，如有请及时清洗，以免影响水位检测精度。',
    'WiFi连接检查：如长时间使用后WiFi连接不稳定，可断电重启设备和手机，重新连接热点。',
    '固件版本检查：关注设备固件版本更新信息，如有新版本可联系技术人员获取最新固件包进行升级。',
    '备用电源：建议使用带过载保护的5V/2A电源适配器，避免因电源不稳定导致设备工作异常。',
]:
    p(item)

# ============================================================
#  附录A 术语表
# ============================================================
page_break()
p('附录A  术语表', style='Heading 1')
p_obj = p('基于物联网的智能宠物喂养控制系统',
          align='center', sb=8, sa=8, indent=False, bold=True, size=16)

terms = [
    ('Android',
     '一种基于Linux内核的开放源代码移动操作系统，由Google公司主导开发，本软件运行于该操作系统之上。'),
    ('Kotlin',
     '一种在Java虚拟机上运行的静态类型编程语言，被Google官方推荐为Android应用开发首选语言，本软件使用Kotlin语言开发。'),
    ('Jetpack Compose',
     'Android官方推荐的现代化声明式UI框架，采用Kotlin语言编写界面，本软件使用Compose构建所有用户界面。'),
    ('TCP/IP',
     '传输控制协议/网际协议，是互联网最基本的通信协议套件。本软件通过TCP协议与设备建立Socket连接进行数据通信。'),
    ('ESP8266',
     '一款低成本、低功耗的WiFi通信模块芯片，支持AP（接入点）和STA（站点）两种工作模式，本系统采用AP模式实现设备与手机的点对点通信。'),
    ('AP模式',
     'Access Point（接入点）模式，即WiFi热点模式。ESP8266工作在AP模式时，自身创建一个WiFi热点，手机直接连接该热点即可与设备通信，无需路由器。'),
    ('Flash存储器',
     '一种电可擦除可编程只读存储器（EEPROM），具有断电数据不丢失的特性。本系统使用STM32内部Flash存储阈值参数，确保断电后参数不丢失。'),
    ('RTC',
     'Real-Time Clock（实时时钟），是一种计算机时钟芯片，能够在系统断电时依靠后备电池继续计时。本系统使用RTC实现定时任务和时间戳功能。'),
    ('步进电机',
     '一种将电脉冲信号转换为角位移的执行机构，每输入一个脉冲信号，电机转动一个固定角度（步距角）。本系统使用28BYJ-48步进电机作为喂食机构的驱动源。'),
    ('HX711',
     '一款专为电子秤设计的24位高精度ADC（模数转换器）芯片，可将称重传感器的模拟信号转换为数字信号。本系统使用HX711采集称重传感器数据，精度可达0.1g。'),
    ('DHT11',
     '一款数字温湿度传感器，采用单总线接口，可同时测量环境温度和相对湿度。本系统使用DHT11采集饲养环境温湿度数据。'),
    ('增量称重',
     '一种称重控制方式，记录喂食开始前的初始重量，以"当前重量-初始重量"作为已喂食重量，当增量达到目标值时停止喂食，可有效消除皮重变化对称重精度的影响。'),
    ('防卡粮',
     '一种步进电机控制策略，当饲料卡住导致电机堵转时，自动反转一定角度使饲料松动，然后再正转继续喂食，循环执行直至喂食完成。'),
]
for term, definition in terms:
    p_obj = doc.add_paragraph()
    pf = p_obj.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.74)
    r1 = p_obj.add_run(term + '：')
    r1.font.bold = True
    r1.font.size = Pt(11)
    set_cjk(r1, '宋体')
    r2 = p_obj.add_run(definition)
    r2.font.size = Pt(11)
    set_cjk(r2, '宋体')

# ============================================================
#  保存
# ============================================================
doc.save(OUTPUT)
print('生成成功：{}'.format(OUTPUT))
print('文件大小：约{}KB'.format(round(os.path.getsize(OUTPUT) / 1024)))
