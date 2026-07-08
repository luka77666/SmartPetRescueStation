import os
import sys
from docx import Document

# 搜索目录
root = r"c:\Users\86157\Desktop\ronghe"

found_files = []

for dirpath, dirnames, filenames in os.walk(root):
    for fn in filenames:
        if fn.lower().endswith(('.docx', '.doc')):
            fp = os.path.join(dirpath, fn)
            try:
                doc = Document(fp)
                has_name = False
                for p in doc.paragraphs:
                    if '罗志杰' in p.text:
                        has_name = True
                        break
                if not has_name:
                    for t in doc.tables:
                        for row in t.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    if '罗志杰' in p.text:
                                        has_name = True
                                        break
                                if has_name:
                                    break
                        if has_name:
                            break
                if has_name:
                    found_files.append(fp)
            except Exception as e:
                pass

print("=" * 60)
print("包含「罗志杰」的Word文档（共{}个）：".format(len(found_files)))
print("=" * 60)
for i, fp in enumerate(found_files, 1):
    # 显示相对路径，方便阅读
    rel = fp.replace(r"c:\Users\86157\Desktop\ronghe", ".")
    print("  [{}] {}".format(i, rel))
print("=" * 60)
print("以上文件需要替换「罗志杰」→「罗智杰」")
print("请确认无误后，我再执行替换。")
